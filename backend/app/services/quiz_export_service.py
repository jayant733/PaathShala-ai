import json
import re
from io import BytesIO

from fastapi import HTTPException

from app.repositories.quiz_repository import QuizRepository


class QuizExportService:
    """Serializes a Quiz into PDF, DOCX, JSON, Google Apps Script, and Google Forms JSON.

    The ``_to_*`` builders are pure functions over an ORM quiz (with questions loaded),
    so they are unit-testable without a database. The async ``to_*_bytes`` methods
    load the quiz from the repository then delegate to the builders.
    """

    def __init__(self, repository: QuizRepository):
        self.repository = repository

    # ------------------------------------------------------------------ loading
    async def _load(self, quiz_id):
        quiz = await self.repository.get_quiz(quiz_id, with_questions=True)
        if not quiz:
            raise HTTPException(status_code=404, detail="Quiz not found")
        return quiz

    # ------------------------------------------------------------------ JSON
    @staticmethod
    def _to_quiz_dict(quiz) -> dict:
        return {
            "id": str(quiz.id),
            "title": quiz.title,
            "description": quiz.description or "",
            "subject": quiz.subject or "",
            "difficulty": quiz.difficulty,
            "duration_minutes": quiz.duration_minutes,
            "number_of_questions": quiz.number_of_questions,
            "status": quiz.status,
            "questions": [
                {
                    "id": str(q.id),
                    "question_text": q.question_text,
                    "question_type": q.question_type,
                    "options": list(q.options or []),
                    "correct_answers": list(q.correct_answers or []),
                    "explanation": q.explanation or "",
                    "difficulty": q.difficulty,
                    "topic": q.topic,
                    "points": q.points,
                    "order_index": q.order_index,
                }
                for q in quiz.questions
            ],
        }

    # ------------------------------------------------------------------ PDF
    @staticmethod
    def _to_pdf_bytes(quiz) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            rightMargin=0.8 * inch, leftMargin=0.8 * inch,
            topMargin=0.8 * inch, bottomMargin=0.8 * inch,
        )
        styles = getSampleStyleSheet()
        story = [Paragraph(quiz.title, styles["Title"])]
        if quiz.description:
            story.append(Paragraph(quiz.description, styles["Normal"]))
        story.append(Paragraph(
            f"Difficulty: {quiz.difficulty}  |  Questions: {quiz.number_of_questions}  |  "
            f"Duration: {quiz.duration_minutes} min",
            styles["Italic"],
        ))
        story.append(Spacer(1, 14))

        letters = "abcdefghij"
        for i, q in enumerate(quiz.questions, start=1):
            story.append(Paragraph(f"<b>{i}. {q.question_text}</b>", styles["BodyText"]))
            if q.question_type in ("MCQ", "multiple"):
                for idx, opt in enumerate(q.options or []):
                    label = letters[idx] if idx < len(letters) else f"{idx + 1}"
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{label}) {opt}", styles["BodyText"]))
            elif q.question_type == "true_false":
                story.append(Paragraph("&nbsp;&nbsp;&nbsp;a) True", styles["BodyText"]))
                story.append(Paragraph("&nbsp;&nbsp;&nbsp;b) False", styles["BodyText"]))
            story.append(Spacer(1, 10))

        story.append(PageBreak())
        story.append(Paragraph("Answer Key", styles["Heading1"]))
        for i, q in enumerate(quiz.questions, start=1):
            answer = ", ".join(q.correct_answers or [])
            story.append(Paragraph(f"<b>{i}.</b> {answer}", styles["BodyText"]))

        doc.build(story)
        return buf.getvalue()

    # ------------------------------------------------------------------ DOCX
    @staticmethod
    def _to_docx_bytes(quiz) -> bytes:
        from docx import Document

        doc = Document()
        doc.add_heading(quiz.title, 0)
        if quiz.description:
            doc.add_paragraph(quiz.description)
        doc.add_paragraph(
            f"Difficulty: {quiz.difficulty}  |  Questions: {quiz.number_of_questions}  |  "
            f"Duration: {quiz.duration_minutes} min"
        )
        for i, q in enumerate(quiz.questions, start=1):
            doc.add_paragraph(f"{i}. {q.question_text}")
            if q.question_type in ("MCQ", "multiple"):
                for opt in (q.options or []):
                    doc.add_paragraph(opt, style="List Bullet")
            elif q.question_type == "true_false":
                doc.add_paragraph("True", style="List Bullet")
                doc.add_paragraph("False", style="List Bullet")
        doc.add_heading("Answer Key", level=1)
        for i, q in enumerate(quiz.questions, start=1):
            doc.add_paragraph(f"{i}. {', '.join(q.correct_answers or [])}")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------ Apps Script
    @staticmethod
    def _to_apps_script(quiz) -> str:
        data = QuizExportService._to_quiz_dict(quiz)
        template = """// Generated by PaathShala AI. Paste into Google Apps Script (Extensions > Apps Script)
// then run createQuizForm() to build a Google Form with the quiz questions and answer key.
const QUIZ_DATA = __QUIZ_JSON__;

function createQuizForm() {
  const form = FormApp.create(QUIZ_DATA.title);
  if (QUIZ_DATA.description) form.setDescription(QUIZ_DATA.description);

  QUIZ_DATA.questions.forEach((q, index) => {
    const title = `${index + 1}. ${q.question_text}`;

    if (q.question_type === 'short_answer') {
      const item = form.addTextItem();
      item.setTitle(title).setRequired(true);
      item.setHelpText(q.explanation || '');
      return;
    }

    if (q.question_type === 'multiple') {
      // CheckboxItem has no server-side grading in Apps Script; choices are added for manual review.
      const item = form.addCheckboxItem();
      item.setTitle(title).setRequired(true);
      item.setChoices(q.options.map(opt => item.createChoice(opt)));
      item.setHelpText(q.explanation || '');
      return;
    }

    const item = form.addMultipleChoiceItem();
    item.setTitle(title).setRequired(true);
    item.setChoices(q.options.map(opt => item.createChoice(opt, isCorrect(opt, q.correct_answers))));
    item.setHelpText(q.explanation || '');
  });

  Logger.log('Form created: %s', form.getPublishedUrl());
  return form.getPublishedUrl();
}

function isCorrect(option, correctAnswers) {
  return correctAnswers.some(ca => String(ca).toLowerCase() === String(option).toLowerCase());
}
"""
        return template.replace("__QUIZ_JSON__", json.dumps(data))

    # ------------------------------------------------------------------ Google Forms JSON
    @staticmethod
    def _to_google_forms_json(quiz) -> dict:
        items = []
        for i, q in enumerate(quiz.questions, start=1):
            question = {
                "questionId": str(q.id),
                "required": True,
                "pointValue": q.points or 1,
            }
            title = f"{i}. {q.question_text}"
            if q.question_type in ("MCQ", "true_false"):
                options = q.options if q.question_type == "MCQ" else ["True", "False"]
                question["choiceQuestion"] = {
                    "type": "RADIO",
                    "options": [{"value": o} for o in options],
                }
                question["grading"] = {
                    "correctAnswers": {"answers": [{"value": ca} for ca in q.correct_answers]}
                }
            elif q.question_type == "multiple":
                question["choiceQuestion"] = {
                    "type": "CHECKBOX",
                    "options": [{"value": o} for o in (q.options or [])],
                }
                question["grading"] = {
                    "correctAnswers": {"answers": [{"value": ca} for ca in q.correct_answers]}
                }
            else:  # short_answer
                question["textQuestion"] = {"paragraph": False}
                question["grading"] = {
                    "correctAnswers": {"answers": [{"value": ca} for ca in q.correct_answers]}
                }
            items.append({"title": title, "questionItem": {"question": question}})

        return {
            "info": {
                "title": quiz.title,
                "description": quiz.description or "",
                "documentTitle": f"{quiz.title} (Quiz)",
            },
            "items": items,
        }

    # ------------------------------------------------------------------ async wrappers (DB-backed)
    async def to_json_bytes(self, quiz_id) -> bytes:
        quiz = await self._load(quiz_id)
        return json.dumps(self._to_quiz_dict(quiz), indent=2).encode("utf-8")

    async def to_pdf_bytes(self, quiz_id) -> bytes:
        quiz = await self._load(quiz_id)
        return self._to_pdf_bytes(quiz)

    async def to_docx_bytes(self, quiz_id) -> bytes:
        quiz = await self._load(quiz_id)
        return self._to_docx_bytes(quiz)

    async def to_apps_script(self, quiz_id) -> str:
        quiz = await self._load(quiz_id)
        return self._to_apps_script(quiz)

    async def to_google_forms_json(self, quiz_id) -> dict:
        quiz = await self._load(quiz_id)
        return self._to_google_forms_json(quiz)

    @staticmethod
    def slug(title: str) -> str:
        return re.sub(r"\W+", "-", title.lower()).strip("-") or "quiz"
